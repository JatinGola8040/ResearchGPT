import io
from datetime import datetime
from typing import Dict, Any, List, Tuple
from docx import Document
from docx.shared import RGBColor
from fpdf import FPDF

HEADING_MAP = {
    "executive_summary": "Executive Summary",
    "key_contributions": "Key Contributions",
    "methodology": "Methodology",
    "results": "Results",
    "limitations": "Limitations",
    "answer": "Answer",
    "research_objective": "Research Objective",
    "research_objectives": "Research Objectives",
    "datasets": "Datasets",
    "datasets_used": "Datasets Used",
    "strengths": "Strengths",
    "key_differences": "Key Differences",
    "overall_conclusion": "Overall Conclusion",
    "conclusion": "Conclusion",
    "current_research_coverage": "Current Research Coverage",
    "common_themes": "Common Themes",
    "conflicting_findings": "Conflicting Findings",
    "research_gaps": "Research Gaps",
    "future_research_opportunities": "Future Research Opportunities",
    "future_scope": "Future Scope",
    "potential_research_questions": "Potential Research Questions",
    "introduction": "Introduction",
    "related_work": "Related Work",
    "methodology_comparison": "Methodology Comparison",
    "key_findings": "Key Findings",
    "research_trends": "Research Trends",
    "papers": "Analyzed Papers",
    "references": "References",
    "citations": "Citations"
}

def sanitize_text_for_pdf(text: str) -> str:
    if not text:
        return ""
    replacements = {
        '“': '"', '”': '"', '‘': "'", '’': "'",
        '—': '--', '–': '-', '•': '*', '…': '...',
        '\u2011': '-', '\u00a0': ' '
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode('latin-1', errors='replace').decode('latin-1')

def extract_report_title(content: Dict[str, Any]) -> str:
    if content.get("title"):
        return str(content["title"])
    if content.get("paper_title"):
        return f"Research Summary: {content['paper_title']}"
    if content.get("comparison"):
        return "Comparative Research Analysis Report"
    if content.get("analysis"):
        return "Research Gap Analysis Report"
    if content.get("answer"):
        return "Research Query Analysis"
    return "ResearchGPT Intelligence Report"

def extract_sections(content: Dict[str, Any]) -> List[Tuple[str, Any]]:
    sections = []
    
    if "papers" in content and isinstance(content["papers"], list) and content["papers"]:
        paper_items = []
        for p in content["papers"]:
            if isinstance(p, dict):
                pid = p.get("paper_id", "")
                ptitle = p.get("paper_title", "Untitled")
                paper_items.append(f"{ptitle} (ID: {pid})" if pid else ptitle)
            else:
                paper_items.append(str(p))
        if paper_items:
            sections.append(("Analyzed Papers", paper_items))

    main_dict = content
    if "comparison" in content and isinstance(content["comparison"], dict):
        main_dict = content["comparison"]
    elif "analysis" in content and isinstance(content["analysis"], dict):
        main_dict = content["analysis"]

    preferred_order = [
        "executive_summary", "introduction", "answer", "research_objective", "research_objectives",
        "current_research_coverage", "related_work", "methodology", "methodology_comparison",
        "datasets", "datasets_used", "key_contributions", "key_findings", "results",
        "strengths", "common_themes", "key_differences", "conflicting_findings",
        "limitations", "research_trends", "research_gaps", "future_research_opportunities",
        "future_scope", "potential_research_questions", "overall_conclusion", "conclusion"
    ]

    processed_keys = {"papers", "title", "paper_title", "paper_id", "mode", "comparison", "analysis", "references", "citations"}
    
    for key in preferred_order:
        if key in main_dict and main_dict[key] is not None:
            val = main_dict[key]
            if val != "" and val != []:
                heading = HEADING_MAP.get(key, key.replace("_", " ").title())
                sections.append((heading, val))
            processed_keys.add(key)

    for d in [main_dict, content]:
        for key, val in d.items():
            if key not in processed_keys and val is not None and val != "" and val != []:
                heading = HEADING_MAP.get(key, key.replace("_", " ").title())
                sections.append((heading, val))
                processed_keys.add(key)

    if "references" in content and isinstance(content["references"], list) and content["references"]:
        ref_items = []
        for r in content["references"]:
            if isinstance(r, dict):
                rtitle = r.get("paper_title", "Reference")
                rcit = r.get("citation", "")
                ref_items.append(f"{rtitle}: {rcit}" if rcit else rtitle)
            else:
                ref_items.append(str(r))
        if ref_items:
            sections.append(("References", ref_items))

    if "citations" in content and isinstance(content["citations"], list) and content["citations"]:
        cit_items = []
        for c in content["citations"]:
            if isinstance(c, dict):
                ctitle = c.get("paper_title", "Paper")
                cpage = c.get("page", "")
                csnip = c.get("snippet", "")
                page_str = f" [Page {cpage}]" if cpage else ""
                snip_str = f' - "{csnip}"' if csnip else ""
                cit_items.append(f"{ctitle}{page_str}{snip_str}")
            else:
                cit_items.append(str(c))
        if cit_items:
            sections.append(("Citations", cit_items))

    return sections

class ResearchPDF(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 15)
        self.set_text_color(41, 128, 185)
        self.cell(0, 10, 'ResearchGPT Intelligence Report', border=0, align='C')
        self.ln(12)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

def generate_pdf_export(content: Dict[str, Any]) -> bytes:
    title = extract_report_title(content)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sections = extract_sections(content)

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        
        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('ReportTitle', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#2980b9'), spaceAfter=10)
        time_style = ParagraphStyle('ReportTime', parent=styles['Italic'], fontSize=10, textColor=colors.HexColor('#6c757d'), spaceAfter=15)
        heading_style = ParagraphStyle('SectionHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2980b9'), spaceBefore=12, spaceAfter=6)
        body_style = ParagraphStyle('SectionBody', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#323232'), spaceAfter=6)
        bullet_style = ParagraphStyle('SectionBullet', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#323232'), leftIndent=15, spaceAfter=4)

        story = []
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Generated on: {timestamp}", time_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#c8c8c8'), spaceAfter=15))

        for heading, val in sections:
            story.append(Paragraph(heading, heading_style))
            if isinstance(val, list):
                for item in val:
                    story.append(Paragraph(f"• {str(item)}", bullet_style))
            else:
                for line in str(val).split('\n'):
                    if line.strip():
                        story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 10))

        doc.build(story)
        bio.seek(0)
        return bio.getvalue()
    except ImportError:
        pdf = ResearchPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font('Helvetica', 'B', 16)
        pdf.set_text_color(33, 37, 41)
        pdf.multi_cell(0, 8, sanitize_text_for_pdf(title))
        pdf.ln(3)
        
        pdf.set_font('Helvetica', 'I', 10)
        pdf.set_text_color(108, 117, 125)
        pdf.cell(0, 6, f'Generated on: {timestamp}', ln=True)
        pdf.ln(8)
        
        for heading, val in sections:
            pdf.set_font('Helvetica', 'B', 13)
            pdf.set_text_color(41, 128, 185)
            pdf.multi_cell(0, 7, sanitize_text_for_pdf(heading))
            pdf.ln(2)
            
            pdf.set_font('Helvetica', '', 11)
            pdf.set_text_color(50, 50, 50)
            
            if isinstance(val, list):
                for item in val:
                    item_str = f" * {str(item)}"
                    pdf.multi_cell(0, 6, sanitize_text_for_pdf(item_str))
                    pdf.ln(1)
            else:
                pdf.multi_cell(0, 6, sanitize_text_for_pdf(str(val)))
            pdf.ln(5)
            
        bio = io.BytesIO()
        pdf.output(bio)
        bio.seek(0)
        return bio.getvalue()

def generate_docx_export(content: Dict[str, Any]) -> bytes:
    title = extract_report_title(content)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sections = extract_sections(content)

    doc = Document()
    doc.add_heading(title, level=0)
    
    p_time = doc.add_paragraph(f"Generated on: {timestamp}")
    if p_time.runs:
        p_time.runs[0].font.italic = True
        p_time.runs[0].font.color.rgb = RGBColor(108, 117, 125)
    
    doc.add_paragraph()
    
    for heading, val in sections:
        doc.add_heading(heading, level=1)
        if isinstance(val, list):
            for item in val:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph(str(val))
        doc.add_paragraph()
        
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
